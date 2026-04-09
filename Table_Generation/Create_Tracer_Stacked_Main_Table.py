import io
import json
import os
import zipfile
from pathlib import Path
from typing import Dict, List, Optional, Sequence

import numpy as np
import pandas as pd

# ============================================================
# USER INPUT
# ============================================================
VIS_SOURCE = r"C:\Users\brend\OneDrive - University of Toronto\VS Code Files\Results\Results_Paper_1\VIS320_Swollen_No_Pressure_135kvp"
GAD_SOURCE = r"C:\Users\brend\OneDrive - University of Toronto\VS Code Files\Results\Results_Paper_1\GAD_Swollen_No_Pressure_135kvp"

# Only set these if both sources point to a shared Results_Paper_1 folder.
VIS_RUN_REL = None
GAD_RUN_REL = None

VIS_ROI_FOLDER = "VIS_320"
GAD_ROI_FOLDER = "GAD"
VIS_LABEL = "VIS 320"
GAD_LABEL = "GAD"

POST_TIME_MIN = 5.0
OUT_FOLDER = r"C:\Users\brend\OneDrive - University of Toronto\VS Code Files\Results\Results_Paper_1\Paper_1_Comparisons\Paper_Tables_Stacked_Main"

WRITE_LATEX = True
WRITE_MARKDOWN = True
WRITE_CSV = True
WRITE_HTML = True
WRITE_DOCX = True

# ============================================================
# FILE HELPERS
# ============================================================
def _is_zip_path(path: str) -> bool:
    return str(path).lower().endswith(".zip")


def _normalize_rel_path(rel_path: str) -> str:
    return str(rel_path).replace("\\", "/").strip("/")


def _list_paths_any(base: str) -> Sequence[str]:
    if _is_zip_path(base):
        with zipfile.ZipFile(base) as zf:
            return zf.namelist()
    base_path = Path(base)
    out = []
    for p in base_path.rglob("*"):
        if p.is_file():
            out.append(str(p.relative_to(base_path)).replace("\\", "/"))
    return out


def _path_exists_any(base: str, relative_path: str) -> bool:
    rel = _normalize_rel_path(relative_path)
    if rel == "":
        return False
    if _is_zip_path(base):
        with zipfile.ZipFile(base) as zf:
            return rel in zf.namelist()
    return (Path(base) / rel).exists()


def _read_csv_any(base: str, relative_path: str) -> pd.DataFrame:
    rel = _normalize_rel_path(relative_path)
    if _is_zip_path(base):
        with zipfile.ZipFile(base) as zf:
            return pd.read_csv(io.BytesIO(zf.read(rel)))
    return pd.read_csv(Path(base) / rel)


def infer_run_prefix(base: str, roi_folder: str) -> str:
    roi_folder = str(roi_folder).strip("/").strip("\\")
    paths = [_normalize_rel_path(p) for p in _list_paths_any(base)]

    root_summary = "multi_roi_summary.csv"
    root_roi_fit = f"{roi_folder}/CSVs_Summaries/fit_parameters_vs_time.csv"
    if root_summary in paths and root_roi_fit in paths:
        return ""

    candidates = []
    suffix = f"/{roi_folder}/CSVs_Summaries/fit_parameters_vs_time.csv"
    for p in paths:
        if p.endswith(suffix):
            prefix = p[:-len(suffix)]
            if prefix and f"{prefix}/multi_roi_summary.csv" in paths:
                candidates.append(prefix)

    candidates = sorted(set(candidates))
    if len(candidates) == 1:
        return candidates[0]
    if len(candidates) > 1:
        raise ValueError(f"Multiple possible run prefixes found in {base}: {candidates}")

    raise FileNotFoundError(
        f"Could not infer the run prefix in {base}. Expected either root-level files or a single run folder containing "
        f"'multi_roi_summary.csv' and '{roi_folder}/CSVs_Summaries/fit_parameters_vs_time.csv'."
    )


def resolve_run_prefix(base: str, explicit_run_rel: Optional[str], roi_folder: str) -> str:
    if explicit_run_rel not in (None, ""):
        return _normalize_rel_path(explicit_run_rel)
    return infer_run_prefix(base, roi_folder)


def join_rel(prefix: str, suffix: str) -> str:
    prefix = _normalize_rel_path(prefix)
    suffix = _normalize_rel_path(suffix)
    return f"{prefix}/{suffix}" if prefix else suffix


def build_tracer_paths(run_prefix: str, roi_folder: str) -> Dict[str, str]:
    roi_folder = str(roi_folder).strip("/").strip("\\")
    roi_base = join_rel(run_prefix, roi_folder)
    return {
        "multi_roi_summary_csv": join_rel(run_prefix, "multi_roi_summary.csv"),
        "fit_parameters_csv": join_rel(roi_base, "CSVs_Summaries/fit_parameters_vs_time.csv"),
        "flux_fixed_roi_std_map_csv": join_rel(roi_base, "CSVs_Uncertainty/diffusive_flux_magnitude_fixed_roi_std_map.csv"),
        "flux_fixed_roi_ci_width_map_csv": join_rel(roi_base, "CSVs_Uncertainty/diffusive_flux_magnitude_fixed_roi_ci_width_map.csv"),
        "flux_combined_std_map_csv": join_rel(roi_base, "CSVs_Uncertainty/diffusive_flux_magnitude_combined_std_map.csv"),
        "flux_combined_ci_width_map_csv": join_rel(roi_base, "CSVs_Uncertainty/diffusive_flux_magnitude_combined_ci_width_map.csv"),
    }


def load_csv_required_any(base: str, rel_path: str) -> pd.DataFrame:
    if not _path_exists_any(base, rel_path):
        raise FileNotFoundError(f"Required file not found: {rel_path}")
    return _read_csv_any(base, rel_path)


def load_csv_optional_any(base: str, rel_path: Optional[str]) -> Optional[pd.DataFrame]:
    if rel_path in (None, ""):
        return None
    if not _path_exists_any(base, rel_path):
        return None
    return _read_csv_any(base, rel_path)


def find_time_column(df: pd.DataFrame) -> str:
    for col in ["time", "time_plot", "time_min", "time_minutes"]:
        if col in df.columns:
            return col
    raise KeyError(f"Could not find a time column in: {df.columns.tolist()}")


def safe_value(df: pd.DataFrame, col: str, row: int = 0):
    if col not in df.columns or len(df) <= row:
        return np.nan
    return df.iloc[row][col]


def find_numeric_series(df: pd.DataFrame, candidates: List[str]) -> Optional[pd.Series]:
    for col in candidates:
        if col in df.columns:
            return pd.to_numeric(df[col], errors="coerce")
    return None


# ============================================================
# FORMATTING HELPERS
# ============================================================
def is_finite_number(x) -> bool:
    try:
        return np.isfinite(float(x))
    except Exception:
        return False


def fmt_fixed(x, decimals=3) -> str:
    if not is_finite_number(x):
        return "NA"
    return f"{float(x):.{decimals}f}"


def fmt_sci_latex(x, sig=3) -> str:
    if not is_finite_number(x):
        return "NA"
    x = float(x)
    if x == 0:
        return "0"
    exponent = int(np.floor(np.log10(abs(x))))
    mantissa = x / (10 ** exponent)
    return rf"${mantissa:.{sig - 1}f} \times 10^{{{exponent}}}$"


def fmt_sci_plain(x, sig=3) -> str:
    if not is_finite_number(x):
        return "NA"
    x = float(x)
    if x == 0:
        return "0"
    exponent = int(np.floor(np.log10(abs(x))))
    mantissa = x / (10 ** exponent)
    return f"{mantissa:.{sig - 1}f} × 10^{exponent}"


def fmt_d_latex(x) -> str:
    return fmt_sci_latex(x, sig=3)


def fmt_d_plain(x) -> str:
    return fmt_sci_plain(x, sig=3)


def fmt_flux(x) -> str:
    if not is_finite_number(x):
        return "NA"
    return f"{float(x):.4f}"


def fmt_r2(x) -> str:
    if not is_finite_number(x):
        return "NA"
    return f"{float(x):.3f}"


def fmt_rmse(x) -> str:
    if not is_finite_number(x):
        return "NA"
    return f"{float(x):.3f}"


def fmt_pm(center, spread, formatter) -> str:
    if not is_finite_number(center) or not is_finite_number(spread):
        return "NA"
    return f"{formatter(center)} ± {formatter(spread)}"


def fmt_median_iqr(median, iqr, formatter) -> str:
    if not is_finite_number(median) or not is_finite_number(iqr):
        return "NA"
    return f"{formatter(median)} [{formatter(iqr)}]"


def fmt_dual_uncertainty(fixed_val, combined_val, formatter) -> str:
    fixed_txt = formatter(fixed_val) if is_finite_number(fixed_val) else "NA"
    combined_txt = formatter(combined_val) if is_finite_number(combined_val) else "NA"
    return f"{fixed_txt} / {combined_txt}"


def maybe_nanmean_rowwise(arr: np.ndarray) -> np.ndarray:
    row_means = np.full(arr.shape[0], np.nan, dtype=float)
    finite_row_mask = np.any(np.isfinite(arr), axis=1)
    if np.any(finite_row_mask):
        row_means[finite_row_mask] = np.nanmean(arr[finite_row_mask], axis=1)
    return row_means


def nanmean_or_nan(arr) -> float:
    arr = np.asarray(arr, dtype=float)
    return float(np.nanmean(arr)) if np.any(np.isfinite(arr)) else np.nan


def latex_escape(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    out = str(text)
    for old, new in replacements.items():
        out = out.replace(old, new)
    return out


# ============================================================
# SUMMARIES
# ============================================================
def summarize_fit_quality(fit_df: pd.DataFrame, post_time_min: float) -> Dict[str, float]:
    tcol = find_time_column(fit_df)
    t = pd.to_numeric(fit_df[tcol], errors="coerce")
    mask_post = t >= float(post_time_min)

    out = {}
    metric_sets = [
        ("profile_fit_r2", "profile_fit_rmse"),
        ("regularized_profile_fit_r2", "regularized_profile_fit_rmse"),
        ("global_profile_fit_r2", "global_profile_fit_rmse"),
    ]
    prefixes = ["per_timepoint", "regularized", "global"]

    for prefix, (r2_col, rmse_col) in zip(prefixes, metric_sets):
        if r2_col in fit_df.columns:
            s = pd.to_numeric(fit_df[r2_col], errors="coerce")
            out[f"{prefix}_r2_mean_all"] = s.mean()
            out[f"{prefix}_r2_mean_post_{int(post_time_min)}min"] = s[mask_post].mean()
        else:
            out[f"{prefix}_r2_mean_all"] = np.nan
            out[f"{prefix}_r2_mean_post_{int(post_time_min)}min"] = np.nan

        if rmse_col in fit_df.columns:
            s = pd.to_numeric(fit_df[rmse_col], errors="coerce")
            out[f"{prefix}_rmse_mean_all"] = s.mean()
            out[f"{prefix}_rmse_mean_post_{int(post_time_min)}min"] = s[mask_post].mean()
        else:
            out[f"{prefix}_rmse_mean_all"] = np.nan
            out[f"{prefix}_rmse_mean_post_{int(post_time_min)}min"] = np.nan

    return out


def summarize_fixed_roi_parameter_uncertainty(fit_df: pd.DataFrame, post_time_min: float) -> Dict[str, float]:
    tcol = find_time_column(fit_df)
    t = pd.to_numeric(fit_df[tcol], errors="coerce")
    mask_post = t >= float(post_time_min)

    d_fit = find_numeric_series(fit_df, [
        "effective_diffusivity_std", "effective_diffusivity_std_mm2_s",
        "D_std_vs_time", "D_std", "fitted_D_std"
    ])
    d_hu = find_numeric_series(fit_df, [
        "effective_diffusivity_hu_noise_std", "effective_diffusivity_hu_noise_std_mm2_s",
        "D_noise_std_vs_time", "D_hu_noise_std"
    ])
    d_cal = find_numeric_series(fit_df, [
        "effective_diffusivity_calibration_std", "effective_diffusivity_calibration_std_mm2_s",
        "D_calibration_std_vs_time", "D_calibration_std"
    ])

    cs_fit = find_numeric_series(fit_df, [
        "fitted_Cs_std", "Cs_std", "Cs_std_vs_time"
    ])
    cs_hu = find_numeric_series(fit_df, [
        "fitted_Cs_hu_noise_std", "Cs_hu_noise_std", "Cs_noise_std_vs_time"
    ])
    cs_cal = find_numeric_series(fit_df, [
        "fitted_Cs_calibration_std", "Cs_calibration_std", "Cs_calibration_std_vs_time"
    ])

    def combine(*series):
        vals = [pd.to_numeric(s, errors="coerce") for s in series if s is not None]
        if not vals:
            return None
        arr = np.vstack([v.to_numpy(dtype=float) for v in vals])
        return pd.Series(np.sqrt(np.nansum(arr ** 2, axis=0)))

    d_fixed = combine(d_fit, d_hu, d_cal)
    cs_fixed = combine(cs_fit, cs_hu, cs_cal)

    out = {
        "D_fixed_roi_std_mean_all": nanmean_or_nan(d_fixed) if d_fixed is not None else np.nan,
        f"D_fixed_roi_std_mean_post_{int(post_time_min)}min": nanmean_or_nan(d_fixed[mask_post]) if d_fixed is not None else np.nan,
        "Cs_fixed_roi_std_mean_all": nanmean_or_nan(cs_fixed) if cs_fixed is not None else np.nan,
        f"Cs_fixed_roi_std_mean_post_{int(post_time_min)}min": nanmean_or_nan(cs_fixed[mask_post]) if cs_fixed is not None else np.nan,
    }
    return out


def summarize_flux_uncertainty_maps(base: str, fit_df: pd.DataFrame, paths: Dict[str, str], post_time_min: float) -> Dict[str, float]:
    tcol = find_time_column(fit_df)
    times = pd.to_numeric(fit_df[tcol], errors="coerce").to_numpy(dtype=float)
    post_mask = times >= float(post_time_min)

    out = {}
    map_specs = {
        "flux_fixed_roi_std": "flux_fixed_roi_std_map_csv",
        "flux_fixed_roi_ci_width": "flux_fixed_roi_ci_width_map_csv",
        "flux_combined_std": "flux_combined_std_map_csv",
        "flux_combined_ci_width": "flux_combined_ci_width_map_csv",
    }

    for out_prefix, key in map_specs.items():
        df = load_csv_optional_any(base, paths.get(key))
        if df is None:
            out[f"{out_prefix}_mean_all"] = np.nan
            out[f"{out_prefix}_mean_post_{int(post_time_min)}min"] = np.nan
            continue

        arr = df.apply(pd.to_numeric, errors="coerce").to_numpy(dtype=float)
        if arr.ndim != 2 or arr.shape[0] != len(times):
            out[f"{out_prefix}_mean_all"] = np.nan
            out[f"{out_prefix}_mean_post_{int(post_time_min)}min"] = np.nan
            continue

        row_means = maybe_nanmean_rowwise(arr)
        out[f"{out_prefix}_mean_all"] = nanmean_or_nan(row_means)
        out[f"{out_prefix}_mean_post_{int(post_time_min)}min"] = nanmean_or_nan(row_means[post_mask])

    return out


def load_tracer_summary(source: str, explicit_run_rel: Optional[str], roi_folder: str, tracer_label: str, post_time_min: float) -> Dict[str, object]:
    run_prefix = resolve_run_prefix(source, explicit_run_rel, roi_folder)
    paths = build_tracer_paths(run_prefix, roi_folder)

    summary_df = load_csv_required_any(source, paths["multi_roi_summary_csv"])
    fit_df = load_csv_required_any(source, paths["fit_parameters_csv"])

    fit_quality = summarize_fit_quality(fit_df, post_time_min)
    param_fixed_roi = summarize_fixed_roi_parameter_uncertainty(fit_df, post_time_min)
    flux_unc = summarize_flux_uncertainty_maps(source, fit_df, paths, post_time_min)

    row = {
        "Tracer": tracer_label,
        "ROI": safe_value(summary_df, "roi_name"),
        "run_prefix": run_prefix,
        "post_time_min": post_time_min,
    }

    summary_cols = [
        "global_fit_D_mm2_s",
        "median_post_5min_D_mm2_s",
        "mean_post_5min_D_mm2_s",
        "std_post_5min_D_mm2_s",
        "iqr_post_5min_D_mm2_s",
        "mean_post_5min_D_combined_std_mm2_s",
        "mean_post_5min_Cs",
        "std_post_5min_Cs",
        "mean_post_5min_Cs_combined_std",
        "mean_diffusive_flux_magnitude_post_5min",
    ]
    for col in summary_cols:
        row[col] = safe_value(summary_df, col)

    row.update(fit_quality)
    row.update(param_fixed_roi)
    row.update(flux_unc)
    return row


# ============================================================
# STACKED TABLE BUILDER
# ============================================================
def build_stacked_table_rows(rows: List[Dict[str, object]]) -> List[Dict[str, str]]:
    out: List[Dict[str, str]] = []
    for row in rows:
        pmin = int(row["post_time_min"])
        d_fixed = row.get(f"D_fixed_roi_std_mean_post_{pmin}min", np.nan)
        d_combined = row.get("mean_post_5min_D_combined_std_mm2_s", np.nan)
        cs_fixed = row.get(f"Cs_fixed_roi_std_mean_post_{pmin}min", np.nan)
        cs_combined = row.get("mean_post_5min_Cs_combined_std", np.nan)
        flux_fixed = row.get(f"flux_fixed_roi_std_mean_post_{pmin}min", np.nan)
        flux_combined = row.get(f"flux_combined_std_mean_post_{pmin}min", np.nan)
        flux_fixed_ci = row.get(f"flux_fixed_roi_ci_width_mean_post_{pmin}min", np.nan)
        flux_combined_ci = row.get(f"flux_combined_ci_width_mean_post_{pmin}min", np.nan)

        out.append({"row_type": "section", "metric": row["Tracer"], "value": ""})
        out.extend([
            {
                "row_type": "data",
                "metric": r"Global fit $D$ (mm$^2$/s)",
                "value": fmt_d_latex(row.get("global_fit_D_mm2_s", np.nan)),
            },
            {
                "row_type": "data",
                "metric": rf"Post-{pmin} min $D$, median [IQR] (mm$^2$/s)",
                "value": fmt_median_iqr(
                    row.get("median_post_5min_D_mm2_s", np.nan),
                    row.get("iqr_post_5min_D_mm2_s", np.nan),
                    fmt_d_plain,
                ),
            },
            {
                "row_type": "data",
                "metric": rf"Post-{pmin} min $D$, mean ± SD (mm$^2$/s)",
                "value": fmt_pm(
                    row.get("mean_post_5min_D_mm2_s", np.nan),
                    row.get("std_post_5min_D_mm2_s", np.nan),
                    fmt_d_plain,
                ),
            },
            {
                "row_type": "data",
                "metric": rf"Post-{pmin} min $D$ fixed-ROI / combined std (mm$^2$/s)",
                "value": fmt_dual_uncertainty(d_fixed, d_combined, fmt_d_plain),
            },
            {
                "row_type": "data",
                "metric": rf"Post-{pmin} min $C_s$, mean ± SD (mg/mL)",
                "value": fmt_pm(
                    row.get("mean_post_5min_Cs", np.nan),
                    row.get("std_post_5min_Cs", np.nan),
                    lambda x: fmt_fixed(x, 2),
                ),
            },
            {
                "row_type": "data",
                "metric": rf"Post-{pmin} min $C_s$ fixed-ROI / combined std (mg/mL)",
                "value": fmt_dual_uncertainty(cs_fixed, cs_combined, lambda x: fmt_fixed(x, 2)),
            },
            {
                "row_type": "data",
                "metric": rf"Post-{pmin} min mean $|J_{{diff}}|$",
                "value": fmt_flux(row.get("mean_diffusive_flux_magnitude_post_5min", np.nan)),
            },
            {
                "row_type": "data",
                "metric": rf"Post-{pmin} min $|J_{{diff}}|$ fixed-ROI / combined std",
                "value": fmt_dual_uncertainty(flux_fixed, flux_combined, fmt_flux),
            },
            {
                "row_type": "data",
                "metric": rf"Post-{pmin} min $|J_{{diff}}|$ fixed-ROI / combined CI width",
                "value": fmt_dual_uncertainty(flux_fixed_ci, flux_combined_ci, fmt_flux),
            },
            {
                "row_type": "data",
                "metric": rf"Post-{pmin} min mean $R^2$",
                "value": fmt_r2(row.get(f"per_timepoint_r2_mean_post_{pmin}min", np.nan)),
            },
            {
                "row_type": "data",
                "metric": rf"Post-{pmin} min mean RMSE",
                "value": fmt_rmse(row.get(f"per_timepoint_rmse_mean_post_{pmin}min", np.nan)),
            },
        ])
    return out


def build_stacked_dataframe(rows: List[Dict[str, object]]) -> pd.DataFrame:
    records = []
    for row in rows:
        pmin = int(row["post_time_min"])
        d_fixed = row.get(f"D_fixed_roi_std_mean_post_{pmin}min", np.nan)
        d_combined = row.get("mean_post_5min_D_combined_std_mm2_s", np.nan)
        cs_fixed = row.get(f"Cs_fixed_roi_std_mean_post_{pmin}min", np.nan)
        cs_combined = row.get("mean_post_5min_Cs_combined_std", np.nan)
        flux_fixed = row.get(f"flux_fixed_roi_std_mean_post_{pmin}min", np.nan)
        flux_combined = row.get(f"flux_combined_std_mean_post_{pmin}min", np.nan)
        flux_fixed_ci = row.get(f"flux_fixed_roi_ci_width_mean_post_{pmin}min", np.nan)
        flux_combined_ci = row.get(f"flux_combined_ci_width_mean_post_{pmin}min", np.nan)

        records.extend([
            {"Tracer": row["Tracer"], "Metric": "Global fit D (mm^2/s)", "Value": fmt_d_plain(row.get("global_fit_D_mm2_s", np.nan))},
            {"Tracer": row["Tracer"], "Metric": f"Post-{pmin} min D, median [IQR] (mm^2/s)", "Value": fmt_median_iqr(row.get("median_post_5min_D_mm2_s", np.nan), row.get("iqr_post_5min_D_mm2_s", np.nan), fmt_d_plain)},
            {"Tracer": row["Tracer"], "Metric": f"Post-{pmin} min D, mean ± SD (mm^2/s)", "Value": fmt_pm(row.get("mean_post_5min_D_mm2_s", np.nan), row.get("std_post_5min_D_mm2_s", np.nan), fmt_d_plain)},
            {"Tracer": row["Tracer"], "Metric": f"Post-{pmin} min D fixed-ROI / combined std (mm^2/s)", "Value": fmt_dual_uncertainty(d_fixed, d_combined, fmt_d_plain)},
            {"Tracer": row["Tracer"], "Metric": f"Post-{pmin} min C_s, mean ± SD (mg/mL)", "Value": fmt_pm(row.get("mean_post_5min_Cs", np.nan), row.get("std_post_5min_Cs", np.nan), lambda x: fmt_fixed(x, 2))},
            {"Tracer": row["Tracer"], "Metric": f"Post-{pmin} min C_s fixed-ROI / combined std (mg/mL)", "Value": fmt_dual_uncertainty(cs_fixed, cs_combined, lambda x: fmt_fixed(x, 2))},
            {"Tracer": row["Tracer"], "Metric": f"Post-{pmin} min mean |Jdiff|", "Value": fmt_flux(row.get("mean_diffusive_flux_magnitude_post_5min", np.nan))},
            {"Tracer": row["Tracer"], "Metric": f"Post-{pmin} min |Jdiff| fixed-ROI / combined std", "Value": fmt_dual_uncertainty(flux_fixed, flux_combined, fmt_flux)},
            {"Tracer": row["Tracer"], "Metric": f"Post-{pmin} min |Jdiff| fixed-ROI / combined CI width", "Value": fmt_dual_uncertainty(flux_fixed_ci, flux_combined_ci, fmt_flux)},
            {"Tracer": row["Tracer"], "Metric": f"Post-{pmin} min mean R^2", "Value": fmt_r2(row.get(f"per_timepoint_r2_mean_post_{pmin}min", np.nan))},
            {"Tracer": row["Tracer"], "Metric": f"Post-{pmin} min mean RMSE", "Value": fmt_rmse(row.get(f"per_timepoint_rmse_mean_post_{pmin}min", np.nan))},
        ])
    return pd.DataFrame(records)


# ============================================================
# WRITERS
# ============================================================
def ensure_dir(path: str):
    os.makedirs(path, exist_ok=True)


def write_latex_table(stacked_rows: List[Dict[str, str]], out_path: Path, caption: str, label: str):
    pieces = [
        r"\documentclass{article}",
        r"\usepackage{booktabs}",
        r"\usepackage[margin=1in]{geometry}",
        r"\usepackage{array}",
        r"\usepackage{caption}",
        r"\begin{document}",
        r"\begin{table}[htbp]",
        r"\centering",
        rf"\caption{{{caption}}}",
        rf"\label{{{label}}}",
        r"\small",
        r"\setlength{\tabcolsep}{8pt}",
        r"\renewcommand{\arraystretch}{1.15}",
        r"\begin{tabular}{p{0.72\linewidth} p{0.22\linewidth}}",
        r"\toprule",
        r"Metric & Value \\",
        r"\midrule",
    ]

    for idx, row in enumerate(stacked_rows):
        if row["row_type"] == "section":
            if idx > 0:
                pieces.append(r"\addlinespace[0.35em]")
            pieces.append(rf"\multicolumn{{2}}{{l}}{{\textbf{{{latex_escape(row['metric'])}}}}} \\")
        else:
            metric = row["metric"]
            value = row["value"] if "$" in row["value"] else latex_escape(row["value"])
            metric_out = metric if "$" in metric else latex_escape(metric)
            pieces.append(f"{metric_out} & {value} \\")

    pieces.extend([
        r"\bottomrule",
        r"\end{tabular}",
        r"\end{table}",
        r"\end{document}",
    ])
    out_path.write_text("\n".join(pieces), encoding="utf-8")


def write_markdown_table(stacked_df: pd.DataFrame, out_path: Path):
    try:
        md = stacked_df.to_markdown(index=False)
    except Exception:
        md = stacked_df.to_string(index=False)
    out_path.write_text(md, encoding="utf-8")


def write_html_table(stacked_rows: List[Dict[str, str]], out_path: Path, title: str):
    body_lines = []
    for row in stacked_rows:
        if row["row_type"] == "section":
            body_lines.append(f'<tr class="section"><td colspan="2">{row["metric"]}</td></tr>')
        else:
            body_lines.append(f'<tr><td>{row["metric"]}</td><td>{row["value"]}</td></tr>')
    html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset=\"utf-8\">
<title>{title}</title>
<style>
body {{ font-family: 'Times New Roman', Times, serif; margin: 30px; }}
.table-wrap {{ max-width: 980px; }}
.caption {{ font-size: 15px; margin-bottom: 8px; }}
table {{ border-collapse: collapse; width: 100%; font-size: 14px; }}
thead th {{ text-align: left; padding: 6px 8px; border-top: 1.6px solid #000; border-bottom: 1px solid #000; }}
tbody td {{ padding: 5px 8px; vertical-align: top; }}
tbody tr:last-child td {{ border-bottom: 1.6px solid #000; }}
.section td {{ padding-top: 10px; font-weight: bold; border-top: 1px solid #777; }}
.value {{ white-space: nowrap; }}
</style>
</head>
<body>
<div class=\"table-wrap\">
<div class=\"caption\"><strong>Table 1.</strong> {title}</div>
<table>
<thead>
<tr><th>Metric</th><th>Value</th></tr>
</thead>
<tbody>
{''.join(body_lines)}
</tbody>
</table>
</div>
</body>
</html>"""
    out_path.write_text(html, encoding="utf-8")


def write_docx_table(stacked_rows: List[Dict[str, str]], out_path: Path, title: str) -> bool:
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt
    except Exception:
        return False

    def set_cell_border(cell, **kwargs):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        for edge in ('top', 'left', 'bottom', 'right'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:' + edge
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)
                for key, value in edge_data.items():
                    element.set(qn('w:' + key), str(value))

    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("Table 1. ")
    run.bold = True
    p.add_run(title)
    p.paragraph_format.space_after = Pt(6)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    for cell in hdr:
        for par in cell.paragraphs:
            for run in par.runs:
                run.bold = True
        set_cell_border(cell, top={"val": "single", "sz": "12", "color": "000000"}, bottom={"val": "single", "sz": "8", "color": "000000"})

    for row in stacked_rows:
        if row["row_type"] == "section":
            cells = table.add_row().cells
            cells[0].merge(cells[1])
            cells[0].text = row["metric"]
            for par in cells[0].paragraphs:
                if par.runs:
                    par.runs[0].bold = True
            set_cell_border(cells[0], top={"val": "single", "sz": "6", "color": "777777"})
        else:
            cells = table.add_row().cells
            cells[0].text = row["metric"].replace("$", "")
            cells[1].text = row["value"].replace("$", "")

    last_row = table.rows[-1].cells
    for cell in last_row:
        set_cell_border(cell, bottom={"val": "single", "sz": "12", "color": "000000"})

    for section in doc.sections:
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    doc.save(out_path)
    return True


# ============================================================
# MAIN
# ============================================================
def main():
    ensure_dir(OUT_FOLDER)

    vis_row = load_tracer_summary(VIS_SOURCE, VIS_RUN_REL, VIS_ROI_FOLDER, VIS_LABEL, POST_TIME_MIN)
    gad_row = load_tracer_summary(GAD_SOURCE, GAD_RUN_REL, GAD_ROI_FOLDER, GAD_LABEL, POST_TIME_MIN)

    rows = [vis_row, gad_row]
    stacked_rows = build_stacked_table_rows(rows)
    stacked_df = build_stacked_dataframe(rows)

    caption = (
        "Summary of late-time transport parameters for VIS 320 and GAD in 1\\% agarose. "
        "Data format depends on the row: global fit $D$; post-5 min $D$ as median [IQR] and mean ± SD; "
        "post-5 min $C_s$ as mean ± SD; and uncertainty comparison rows reported as fixed-ROI / combined late-time means."
    )

    if WRITE_CSV:
        stacked_df.to_csv(Path(OUT_FOLDER) / "Table_1_stacked_main_results_v2.csv", index=False)
    if WRITE_MARKDOWN:
        write_markdown_table(stacked_df, Path(OUT_FOLDER) / "Table_1_stacked_main_results_v2.md")
    if WRITE_HTML:
        write_html_table(stacked_rows, Path(OUT_FOLDER) / "Table_1_stacked_main_results_v2.html", title="Summary of late-time transport parameters for VIS 320 and GAD in 1% agarose.")
    if WRITE_LATEX:
        write_latex_table(
            stacked_rows,
            Path(OUT_FOLDER) / "Table_1_stacked_main_results_v2.tex",
            caption=caption,
            label="tab:stacked_main_results_v2",
        )

    wrote_docx = False
    if WRITE_DOCX:
        wrote_docx = write_docx_table(
            stacked_rows,
            Path(OUT_FOLDER) / "Table_1_stacked_main_results_v2.docx",
            title="Summary of late-time transport parameters for VIS 320 and GAD in 1% agarose.",
        )

    audit = {
        "vis_source": VIS_SOURCE,
        "gad_source": GAD_SOURCE,
        "vis_run_prefix": vis_row["run_prefix"],
        "gad_run_prefix": gad_row["run_prefix"],
        "post_time_min": POST_TIME_MIN,
        "wrote_csv": WRITE_CSV,
        "wrote_markdown": WRITE_MARKDOWN,
        "wrote_html": WRITE_HTML,
        "wrote_latex": WRITE_LATEX,
        "wrote_docx": wrote_docx,
        "table_files": [
            "Table_1_stacked_main_results_v2.csv",
            "Table_1_stacked_main_results_v2.md",
            "Table_1_stacked_main_results_v2.html",
            "Table_1_stacked_main_results_v2.tex",
        ] + (["Table_1_stacked_main_results_v2.docx"] if wrote_docx else []),
    }
    (Path(OUT_FOLDER) / "Table_1_stacked_main_results_v2_audit.json").write_text(json.dumps(audit, indent=2), encoding="utf-8")

    print("Saved stacked main-results table files to:", OUT_FOLDER)
    for name in audit["table_files"]:
        print(" -", name)


if __name__ == "__main__":
    main()
