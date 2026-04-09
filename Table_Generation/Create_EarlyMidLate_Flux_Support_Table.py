import io
import os
import json
import zipfile
from pathlib import Path
from typing import Optional, Sequence, Dict, Tuple, List

import numpy as np
import pandas as pd

# ============================================================
# USER INPUT
# ============================================================
# Point each source to either:
#   1) a tracer-specific run folder / zip containing multi_roi_summary.csv and ROI folders, or
#   2) a shared Results_Paper_1-style root, together with VIS_RUN_REL / GAD_RUN_REL below.
VIS_SOURCE = r"C:\Users\brend\OneDrive - University of Toronto\VS Code Files\Results\Results_Paper_1\VIS320_Swollen_No_Pressure_135kvp"
GAD_SOURCE = r"C:\Users\brend\OneDrive - University of Toronto\VS Code Files\Results\Results_Paper_1\GAD_Swollen_No_Pressure_135kvp"

# Only set these if the source points to a shared Results_Paper_1 root instead of a single run folder.
VIS_RUN_REL = None
GAD_RUN_REL = None

VIS_ROI_FOLDER = "VIS_320"
GAD_ROI_FOLDER = "GAD"
VIS_LABEL = "VIS 320"
GAD_LABEL = "GAD"

# Depth spacing used to convert row index -> depth in mm.
VIS_DX_MM = 0.166
GAD_DX_MM = 0.166

# Set to None to auto-pick overlap start / mid / end from the two time vectors.
MANUAL_TARGET_TIMES_MIN = [0.75, 11.0, 21.5]
PANEL_NAMES = ["Beginning", "Middle", "End"]

OUT_FOLDER = r"C:\Users\brend\OneDrive - University of Toronto\VS Code Files\Results\Results_Paper_1\Paper_1_Comparisons\EarlyMidLate_Flux_Support_Table"

WRITE_DOCX = True
WRITE_LATEX = True
WRITE_MARKDOWN = True
WRITE_CSV = True
WRITE_HTML = True

# Relative uncertainty helper
RELATIVE_UNCERTAINTY_EPS = 1e-12
Z95 = 1.96

# Column label helper
FLUX_UNITS = "mg/(mm^2 s)"

# ============================================================
# FILE HELPERS
# ============================================================
def _is_zip_path(path: str) -> bool:
    return str(path).lower().endswith(".zip")


def _normalize_rel_path(rel_path: str) -> str:
    return str(rel_path).replace("\\", "/").strip("/")


def _list_paths_any(base: str) -> Sequence[str]:
    if _is_zip_path(base):
        with zipfile.ZipFile(base) as zf:
            return zf.namelist()
    base_path = Path(base)
    out = []
    for p in base_path.rglob("*"):
        if p.is_file():
            out.append(str(p.relative_to(base_path)).replace("\\", "/"))
    return out


def _path_exists_any(base: str, relative_path: str) -> bool:
    rel = _normalize_rel_path(relative_path)
    if rel == "":
        return False
    if _is_zip_path(base):
        with zipfile.ZipFile(base) as zf:
            return rel in zf.namelist()
    return (Path(base) / rel).exists()


def _read_csv_any(base: str, relative_path: str) -> pd.DataFrame:
    rel = _normalize_rel_path(relative_path)
    if _is_zip_path(base):
        with zipfile.ZipFile(base) as zf:
            return pd.read_csv(io.BytesIO(zf.read(rel)))
    return pd.read_csv(Path(base) / rel)


def infer_run_prefix(base: str, roi_folder: str) -> str:
    """
    Returns:
      ""          -> tracer zip/folder has files directly at root, e.g. multi_roi_summary.csv and ROI folders
      "RUN_NAME"  -> shared Results_Paper_1-style root
    """
    roi_folder = str(roi_folder).strip("/").strip("\\")
    paths = [_normalize_rel_path(p) for p in _list_paths_any(base)]

    root_summary = "multi_roi_summary.csv"
    root_roi_fit = f"{roi_folder}/CSVs_Summaries/fit_parameters_vs_time.csv"
    if root_summary in paths and root_roi_fit in paths:
        return ""

    candidates = []
    suffix = f"/{roi_folder}/CSVs_Summaries/fit_parameters_vs_time.csv"
    for p in paths:
        if p.endswith(suffix):
            prefix = p[:-len(suffix)]
            if prefix and f"{prefix}/multi_roi_summary.csv" in paths:
                candidates.append(prefix)

    candidates = sorted(set(candidates))
    if len(candidates) == 1:
        return candidates[0]
    if len(candidates) > 1:
        raise ValueError(f"Multiple possible run prefixes found in {base}: {candidates}")

    raise FileNotFoundError(
        f"Could not infer the run prefix in {base}. Expected either root-level files or a single run folder containing "
        f"'multi_roi_summary.csv' and '{roi_folder}/CSVs_Summaries/fit_parameters_vs_time.csv'."
    )


def resolve_run_prefix(base: str, explicit_run_rel: Optional[str], roi_folder: str) -> str:
    if explicit_run_rel not in (None, ""):
        return _normalize_rel_path(explicit_run_rel)
    return infer_run_prefix(base, roi_folder)


def join_rel(prefix: str, suffix: str) -> str:
    prefix = _normalize_rel_path(prefix)
    suffix = _normalize_rel_path(suffix)
    return f"{prefix}/{suffix}" if prefix else suffix


def build_tracer_paths(run_prefix: str, roi_folder: str) -> Dict[str, str]:
    roi_folder = str(roi_folder).strip("/").strip("\\")
    roi_base = join_rel(run_prefix, roi_folder)
    return {
        "fit_parameters_csv": join_rel(roi_base, "CSVs_Summaries/fit_parameters_vs_time.csv"),
        "fitted_profiles_csv": join_rel(roi_base, "CSVs_Profiles/fitted_profiles_depth_vs_time.csv"),
        "flux_map_csv": join_rel(roi_base, "CSVs_Diffusion/diffusive_flux_magnitude_map.csv"),
        "fit_std_csv": join_rel(roi_base, "CSVs_Uncertainty/fitted_profiles_std_depth_vs_time.csv"),
        "hu_noise_std_csv": join_rel(roi_base, "CSVs_Uncertainty/fitted_profiles_hu_noise_std_depth_vs_time.csv"),
        "calibration_std_csv": join_rel(roi_base, "CSVs_Uncertainty/fitted_profiles_calibration_std_depth_vs_time.csv"),
    }


def load_csv_required_any(base: str, rel_path: str) -> pd.DataFrame:
    if not _path_exists_any(base, rel_path):
        raise FileNotFoundError(f"Required file not found: {rel_path}")
    return _read_csv_any(base, rel_path)


def load_csv_optional_any(base: str, rel_path: Optional[str]) -> Optional[pd.DataFrame]:
    if rel_path in (None, ""):
        return None
    if not _path_exists_any(base, rel_path):
        return None
    return _read_csv_any(base, rel_path)


# ============================================================
# DATA HELPERS
# ============================================================
def find_time_column(df: pd.DataFrame) -> str:
    for col in ["time", "time_plot", "time_min", "time_minutes"]:
        if col in df.columns:
            return col
    raise KeyError(f"Could not find a time column in: {df.columns.tolist()}")


def to_numeric_2d(df: pd.DataFrame) -> np.ndarray:
    arr = df.apply(pd.to_numeric, errors="coerce").to_numpy(dtype=float)
    if arr.ndim != 2:
        raise ValueError("Expected a 2D numeric array.")
    return arr


def build_depth_mm(n_depth: int, dx_mm: float) -> np.ndarray:
    return np.arange(n_depth, dtype=float) * float(dx_mm)


def choose_target_times(vis_times: np.ndarray, gad_times: np.ndarray) -> List[float]:
    if MANUAL_TARGET_TIMES_MIN is not None:
        if len(MANUAL_TARGET_TIMES_MIN) != 3:
            raise ValueError("MANUAL_TARGET_TIMES_MIN must contain exactly 3 values.")
        return [float(x) for x in MANUAL_TARGET_TIMES_MIN]

    vis_valid = vis_times[np.isfinite(vis_times)]
    gad_valid = gad_times[np.isfinite(gad_times)]
    overlap_start = max(float(np.min(vis_valid)), float(np.min(gad_valid)))
    overlap_end = min(float(np.max(vis_valid)), float(np.max(gad_valid)))
    return [overlap_start, 0.5 * (overlap_start + overlap_end), overlap_end]


def nearest_index(values: np.ndarray, target: float, valid_mask: Optional[np.ndarray] = None) -> int:
    vals = np.asarray(values, dtype=float)
    if valid_mask is None:
        valid_mask = np.isfinite(vals)
    if not np.any(valid_mask):
        raise ValueError("No valid values available for nearest-index selection.")
    idx_valid = np.where(valid_mask)[0]
    nearest_pos = np.argmin(np.abs(vals[idx_valid] - float(target)))
    return int(idx_valid[nearest_pos])


def combine_available_std_maps(*maps: Optional[np.ndarray]) -> Optional[np.ndarray]:
    valid_maps = [np.asarray(m, dtype=float) for m in maps if m is not None]
    if not valid_maps:
        return None
    ref_shape = valid_maps[0].shape
    for m in valid_maps[1:]:
        if m.shape != ref_shape:
            raise ValueError(f"Uncertainty map shape mismatch: expected {ref_shape}, got {m.shape}")
    stack = np.stack(valid_maps, axis=0)
    finite_mask = np.isfinite(stack)
    any_finite = np.any(finite_mask, axis=0)
    rss = np.sqrt(np.nansum(np.where(finite_mask, stack, 0.0) ** 2, axis=0))
    rss[~any_finite] = np.nan
    return rss


def get_r2_column(df: pd.DataFrame) -> Optional[str]:
    for col in ["profile_fit_r2", "r2", "R2"]:
        if col in df.columns:
            return col
    return None


def split_depth_bands(n_depth: int) -> Dict[str, Tuple[int, int]]:
    edges = np.linspace(0, n_depth, 4, dtype=int)
    out = {}
    labels = ["shallow", "mid", "deep"]
    for name, start, end in zip(labels, edges[:-1], edges[1:]):
        out[name] = (int(start), int(end))
    return out


def mean_over_slice(vec: np.ndarray, start: int, end: int) -> float:
    part = np.asarray(vec[start:end], dtype=float)
    return float(np.nanmean(part)) if np.any(np.isfinite(part)) else np.nan


def compute_relative_ci_width_percent(ci_width: np.ndarray,
                                      center_values: np.ndarray,
                                      eps: float = RELATIVE_UNCERTAINTY_EPS) -> np.ndarray:
    ci_width = np.asarray(ci_width, dtype=float)
    center_values = np.asarray(center_values, dtype=float)
    out = np.full_like(ci_width, np.nan, dtype=float)
    valid = np.isfinite(ci_width) & np.isfinite(center_values)
    if np.any(valid):
        out[valid] = 100.0 * ci_width[valid] / np.maximum(np.abs(center_values[valid]), float(eps))
    return out


def get_first_existing_numeric_column(df: pd.DataFrame, candidates: List[str], default: Optional[float] = None) -> np.ndarray:
    for col in candidates:
        if col in df.columns:
            return pd.to_numeric(df[col], errors="coerce").to_numpy(dtype=float)
    if default is None:
        raise KeyError(f"None of the candidate columns were found: {candidates}")
    return np.full(len(df), float(default), dtype=float)


def combine_std_series(*series: Optional[np.ndarray]) -> Optional[np.ndarray]:
    valid = [np.asarray(s, dtype=float) for s in series if s is not None]
    if not valid:
        return None
    ref_len = len(valid[0])
    for s in valid[1:]:
        if len(s) != ref_len:
            raise ValueError(f"Uncertainty series length mismatch: expected {ref_len}, got {len(s)}")
    stack = np.stack(valid, axis=0)
    finite_mask = np.isfinite(stack)
    any_finite = np.any(finite_mask, axis=0)
    rss = np.sqrt(np.nansum(np.where(finite_mask, stack, 0.0) ** 2, axis=0))
    rss[~any_finite] = np.nan
    return rss


def compute_gradient_map(depth_profiles: np.ndarray, dx_mm: float) -> np.ndarray:
    return np.gradient(np.asarray(depth_profiles, dtype=float), float(dx_mm), axis=1)


def propagate_flux_fixed_roi_ci_map(flux_center: np.ndarray,
                                    fitted_profiles_center: np.ndarray,
                                    dx_mm: float,
                                    d_center: np.ndarray,
                                    fitted_profiles_fixed_std: Optional[np.ndarray],
                                    d_fixed_std: Optional[np.ndarray],
                                    z: float = Z95) -> Tuple[Optional[np.ndarray], Optional[np.ndarray], Optional[np.ndarray]]:
    """
    Approximate fixed-ROI 95% CI for |J_diff| = | -D * dC/dx | by propagating
    fixed-ROI uncertainty in fitted concentration profiles and fitted effective diffusivity.

    The bounds are estimated from all low/high combinations of D and C.
    """
    if fitted_profiles_fixed_std is None and d_fixed_std is None:
        return None, None, None

    c_center = np.asarray(fitted_profiles_center, dtype=float)
    j_center = np.asarray(flux_center, dtype=float)
    d_center = np.asarray(d_center, dtype=float)

    if c_center.shape != j_center.shape:
        raise ValueError(f"Flux center map shape {j_center.shape} does not match fitted profile map shape {c_center.shape}")
    if d_center.ndim != 1 or d_center.shape[0] != c_center.shape[0]:
        raise ValueError(f"D center series must have length {c_center.shape[0]}, got {d_center.shape}")

    if fitted_profiles_fixed_std is None:
        c_low = c_center.copy()
        c_high = c_center.copy()
    else:
        c_std = np.asarray(fitted_profiles_fixed_std, dtype=float)
        if c_std.shape != c_center.shape:
            raise ValueError(f"Fixed ROI fitted-profile std map shape {c_std.shape} does not match fitted profile map shape {c_center.shape}")
        c_low = np.maximum(c_center - z * c_std, 0.0)
        c_high = np.maximum(c_center + z * c_std, 0.0)

    if d_fixed_std is None:
        d_low = np.maximum(d_center, 0.0)
        d_high = np.maximum(d_center, 0.0)
    else:
        d_std = np.asarray(d_fixed_std, dtype=float)
        if d_std.ndim != 1 or d_std.shape[0] != c_center.shape[0]:
            raise ValueError(f"Fixed ROI D std series must have length {c_center.shape[0]}, got {d_std.shape}")
        d_low = np.maximum(d_center - z * d_std, 0.0)
        d_high = np.maximum(d_center + z * d_std, 0.0)

    grad_low = compute_gradient_map(c_low, dx_mm)
    grad_high = compute_gradient_map(c_high, dx_mm)

    flux_candidates = np.stack([
        np.abs(-d_low[:, None] * grad_low),
        np.abs(-d_low[:, None] * grad_high),
        np.abs(-d_high[:, None] * grad_low),
        np.abs(-d_high[:, None] * grad_high),
    ], axis=0)

    flux_ci_low = np.nanmin(flux_candidates, axis=0)
    flux_ci_high = np.nanmax(flux_candidates, axis=0)
    flux_ci_width = flux_ci_high - flux_ci_low

    # Keep NaNs aligned with the center map where the flux itself is unavailable.
    invalid_center = ~np.isfinite(j_center)
    flux_ci_low[invalid_center] = np.nan
    flux_ci_high[invalid_center] = np.nan
    flux_ci_width[invalid_center] = np.nan
    return flux_ci_low, flux_ci_high, flux_ci_width


# ============================================================
# FORMATTING HELPERS
# ============================================================
def ensure_dir(path: str):
    os.makedirs(path, exist_ok=True)


def is_finite_number(x) -> bool:
    try:
        return np.isfinite(float(x))
    except Exception:
        return False


def fmt_fixed(x: float, decimals: int = 2) -> str:
    if not is_finite_number(x):
        return "NA"
    return f"{float(x):.{decimals}f}"


def latex_escape(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    out = str(text)
    for old, new in replacements.items():
        out = out.replace(old, new)
    return out


# ============================================================
# CORE TABLE BUILDERS
# ============================================================
def load_tracer_payload(source: str, run_rel: Optional[str], roi_folder: str, tracer_label: str, dx_mm: float) -> Dict[str, object]:
    run_prefix = resolve_run_prefix(source, run_rel, roi_folder)
    paths = build_tracer_paths(run_prefix, roi_folder)

    fit_df = load_csv_required_any(source, paths["fit_parameters_csv"])
    fitted_profiles = to_numeric_2d(load_csv_required_any(source, paths["fitted_profiles_csv"]))
    flux_map = to_numeric_2d(load_csv_required_any(source, paths["flux_map_csv"]))

    fit_std_df = load_csv_optional_any(source, paths["fit_std_csv"])
    hu_std_df = load_csv_optional_any(source, paths["hu_noise_std_csv"])
    cal_std_df = load_csv_optional_any(source, paths["calibration_std_csv"])

    fit_std = to_numeric_2d(fit_std_df) if fit_std_df is not None else None
    hu_std = to_numeric_2d(hu_std_df) if hu_std_df is not None else None
    cal_std = to_numeric_2d(cal_std_df) if cal_std_df is not None else None
    fixed_roi_profile_std = combine_available_std_maps(fit_std, hu_std, cal_std)

    tcol = find_time_column(fit_df)
    times = pd.to_numeric(fit_df[tcol], errors="coerce").to_numpy(dtype=float)
    r2_col = get_r2_column(fit_df)
    r2 = pd.to_numeric(fit_df[r2_col], errors="coerce").to_numpy(dtype=float) if r2_col is not None else np.full(len(times), np.nan, dtype=float)

    d_center = get_first_existing_numeric_column(
        fit_df,
        ["effective_diffusivity", "effective_diffusivity_mm2_s"],
    )
    d_model_std = get_first_existing_numeric_column(
        fit_df,
        ["effective_diffusivity_std", "effective_diffusivity_std_mm2_s"],
        default=np.nan,
    )
    d_hu_std = get_first_existing_numeric_column(
        fit_df,
        ["effective_diffusivity_hu_noise_std", "effective_diffusivity_hu_noise_std_mm2_s"],
        default=np.nan,
    )
    d_cal_std = get_first_existing_numeric_column(
        fit_df,
        ["effective_diffusivity_calibration_std", "effective_diffusivity_calibration_std_mm2_s"],
        default=np.nan,
    )
    d_fixed_roi_std = combine_std_series(d_model_std, d_hu_std, d_cal_std)

    if fitted_profiles.shape[0] != len(times):
        raise ValueError(
            f"{tracer_label}: fitted_profiles has {fitted_profiles.shape[0]} rows but fit_parameters has {len(times)} timepoints."
        )
    if flux_map.shape != fitted_profiles.shape:
        raise ValueError(
            f"{tracer_label}: flux map shape {flux_map.shape} does not match fitted profile map shape {fitted_profiles.shape}."
        )
    if len(d_center) != len(times):
        raise ValueError(
            f"{tracer_label}: effective diffusivity series has {len(d_center)} values but expected {len(times)}."
        )
    if fixed_roi_profile_std is not None and fixed_roi_profile_std.shape != fitted_profiles.shape:
        raise ValueError(
            f"{tracer_label}: fixed ROI fitted-profile std map shape {fixed_roi_profile_std.shape} does not match fitted profile map {fitted_profiles.shape}."
        )

    flux_ci_low, flux_ci_high, flux_ci_width = propagate_flux_fixed_roi_ci_map(
        flux_center=flux_map,
        fitted_profiles_center=fitted_profiles,
        dx_mm=float(dx_mm),
        d_center=d_center,
        fitted_profiles_fixed_std=fixed_roi_profile_std,
        d_fixed_std=d_fixed_roi_std,
        z=Z95,
    )

    return {
        "Tracer": tracer_label,
        "run_prefix": run_prefix,
        "fit_df": fit_df,
        "times": times,
        "flux_map": flux_map,
        "flux_ci_low": flux_ci_low,
        "flux_ci_high": flux_ci_high,
        "flux_ci_width": flux_ci_width,
        "depth_mm": build_depth_mm(fitted_profiles.shape[1], dx_mm),
        "r2": r2,
        "d_center": d_center,
        "d_fixed_roi_std": d_fixed_roi_std,
        "fixed_roi_profile_std": fixed_roi_profile_std,
    }


def build_support_rows(vis: Dict[str, object], gad: Dict[str, object]) -> pd.DataFrame:
    target_times = choose_target_times(vis["times"], gad["times"])

    rows_out: List[Dict[str, object]] = []
    for payload in [vis, gad]:
        tracer = payload["Tracer"]
        times = np.asarray(payload["times"], dtype=float)
        flux_map = np.asarray(payload["flux_map"], dtype=float)
        flux_ci_width = payload["flux_ci_width"]
        r2 = np.asarray(payload["r2"], dtype=float)
        band_idx = split_depth_bands(flux_map.shape[1])

        valid_mask = np.any(np.isfinite(flux_map), axis=1)
        for panel_name, target_t in zip(PANEL_NAMES, target_times):
            idx = nearest_index(times, target_t, valid_mask)
            flux_vec = flux_map[idx]
            ci_width_vec = None if flux_ci_width is None else np.asarray(flux_ci_width[idx], dtype=float)
            relative_ci_percent = None if ci_width_vec is None else compute_relative_ci_width_percent(ci_width_vec, flux_vec)

            row = {
                "Tracer": tracer,
                "Panel": panel_name,
                "Target time (min)": float(target_t),
                "Actual time (min)": float(times[idx]),
                f"Shallow mean diffusive flux magnitude ({FLUX_UNITS})": mean_over_slice(flux_vec, *band_idx["shallow"]),
                f"Mid mean diffusive flux magnitude ({FLUX_UNITS})": mean_over_slice(flux_vec, *band_idx["mid"]),
                f"Deep mean diffusive flux magnitude ({FLUX_UNITS})": mean_over_slice(flux_vec, *band_idx["deep"]),
                f"Overall mean diffusive flux magnitude ({FLUX_UNITS})": mean_over_slice(flux_vec, 0, flux_map.shape[1]),
                f"Mean fixed-ROI 95% CI width ({FLUX_UNITS})": mean_over_slice(ci_width_vec, 0, flux_map.shape[1]) if ci_width_vec is not None else np.nan,
                "Shallow mean fixed-ROI relative 95% CI width (%)": mean_over_slice(relative_ci_percent, *band_idx["shallow"]) if relative_ci_percent is not None else np.nan,
                "Mid mean fixed-ROI relative 95% CI width (%)": mean_over_slice(relative_ci_percent, *band_idx["mid"]) if relative_ci_percent is not None else np.nan,
                "Deep mean fixed-ROI relative 95% CI width (%)": mean_over_slice(relative_ci_percent, *band_idx["deep"]) if relative_ci_percent is not None else np.nan,
                "Profile fit R²": float(r2[idx]) if idx < len(r2) else np.nan,
            }
            rows_out.append(row)

    return pd.DataFrame(rows_out)


def build_display_df(raw_df: pd.DataFrame) -> pd.DataFrame:
    shallow_col = f"Shallow mean diffusive flux magnitude ({FLUX_UNITS})"
    mid_col = f"Mid mean diffusive flux magnitude ({FLUX_UNITS})"
    deep_col = f"Deep mean diffusive flux magnitude ({FLUX_UNITS})"
    overall_col = f"Overall mean diffusive flux magnitude ({FLUX_UNITS})"
    ci_width_col = f"Mean fixed-ROI 95% CI width ({FLUX_UNITS})"

    out_rows = []
    for tracer in [VIS_LABEL, GAD_LABEL]:
        out_rows.append({
            "Metric": tracer,
            "Target time (min)": "",
            "Actual time (min)": "",
            shallow_col: "",
            mid_col: "",
            deep_col: "",
            overall_col: "",
            ci_width_col: "",
            "Shallow mean fixed-ROI relative 95% CI width (%)": "",
            "Mid mean fixed-ROI relative 95% CI width (%)": "",
            "Deep mean fixed-ROI relative 95% CI width (%)": "",
            "Profile fit R²": "",
        })
        sub = raw_df[raw_df["Tracer"] == tracer].copy()
        panel_order = {name: i for i, name in enumerate(PANEL_NAMES)}
        sub["_panel_order"] = sub["Panel"].map(panel_order)
        sub = sub.sort_values(["_panel_order", "Actual time (min)"])
        for _, row in sub.iterrows():
            out_rows.append({
                "Metric": row["Panel"],
                "Target time (min)": fmt_fixed(row["Target time (min)"], 2),
                "Actual time (min)": fmt_fixed(row["Actual time (min)"], 2),
                shallow_col: fmt_fixed(row[shallow_col], 4),
                mid_col: fmt_fixed(row[mid_col], 4),
                deep_col: fmt_fixed(row[deep_col], 4),
                overall_col: fmt_fixed(row[overall_col], 4),
                ci_width_col: fmt_fixed(row[ci_width_col], 4),
                "Shallow mean fixed-ROI relative 95% CI width (%)": fmt_fixed(row["Shallow mean fixed-ROI relative 95% CI width (%)"], 2),
                "Mid mean fixed-ROI relative 95% CI width (%)": fmt_fixed(row["Mid mean fixed-ROI relative 95% CI width (%)"], 2),
                "Deep mean fixed-ROI relative 95% CI width (%)": fmt_fixed(row["Deep mean fixed-ROI relative 95% CI width (%)"], 2),
                "Profile fit R²": fmt_fixed(row["Profile fit R²"], 3),
            })
    return pd.DataFrame(out_rows)


# ============================================================
# OUTPUT WRITERS
# ============================================================
def save_tables_plain(df: pd.DataFrame, out_base: Path):
    if WRITE_CSV:
        df.to_csv(out_base.with_suffix(".csv"), index=False)
    if WRITE_MARKDOWN:
        try:
            md = df.to_markdown(index=False)
        except Exception:
            md = df.to_string(index=False)
        out_base.with_suffix(".md").write_text(md, encoding="utf-8")
    if WRITE_HTML:
        html_rows = []
        for _, row in df.iterrows():
            first_col = str(row.iloc[0])
            if first_col in [VIS_LABEL, GAD_LABEL]:
                html_rows.append(f'<tr class="section"><td colspan="{len(df.columns)}">{first_col}</td></tr>')
            else:
                tds = "".join(f"<td>{row[col]}</td>" for col in df.columns)
                html_rows.append(f"<tr>{tds}</tr>")

        html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>Early / mid / late diffusive flux support table</title>
<style>
body {{ font-family: 'Times New Roman', Times, serif; margin: 30px; }}
.table-wrap {{ max-width: 1800px; }}
.caption {{ font-size: 15px; margin-bottom: 8px; }}
table {{ border-collapse: collapse; width: 100%; font-size: 13px; }}
thead th {{ text-align: left; padding: 6px 8px; border-top: 1.6px solid #000; border-bottom: 1px solid #000; }}
tbody td {{ padding: 5px 8px; vertical-align: top; }}
tbody tr:last-child td {{ border-bottom: 1.6px solid #000; }}
.section td {{ padding-top: 10px; font-weight: bold; border-top: 1px solid #777; }}
</style>
</head>
<body>
<div class="table-wrap">
<div class="caption"><strong>Table 1.</strong> Early-, mid-, and late-time diffusive flux magnitude summaries for VIS 320 and GAD across shallow, mid, and deep depth bands, with propagated fixed-ROI absolute and relative flux uncertainty.</div>
<table>
<thead>
<tr>{''.join(f'<th>{col}</th>' for col in df.columns)}</tr>
</thead>
<tbody>
{''.join(html_rows)}
</tbody>
</table>
</div>
</body>
</html>"""
        out_base.with_suffix(".html").write_text(html, encoding="utf-8")


def df_to_latex_booktabs_stacked(raw_df: pd.DataFrame, caption: str, label: str) -> str:
    shallow_col = f"Shallow mean diffusive flux magnitude ({FLUX_UNITS})"
    mid_col = f"Mid mean diffusive flux magnitude ({FLUX_UNITS})"
    deep_col = f"Deep mean diffusive flux magnitude ({FLUX_UNITS})"
    overall_col = f"Overall mean diffusive flux magnitude ({FLUX_UNITS})"
    ci_width_col = f"Mean fixed-ROI 95% CI width ({FLUX_UNITS})"

    lines = [
        r"\begin{table}[htbp]",
        r"\centering",
        r"\footnotesize",
        r"\setlength{\tabcolsep}{4pt}",
        r"\renewcommand{\arraystretch}{1.08}",
        rf"\caption{{{caption}}}",
        rf"\label{{{label}}}",
        r"\resizebox{\linewidth}{!}{%",
        r"\begin{tabular}{@{}llllllllllll@{}}",
        r"\toprule",
        r"\textbf{Metric} & \shortstack[l]{\textbf{Target time} \\ \textbf{(min)}} & \shortstack[l]{\textbf{Actual time} \\ \textbf{(min)}} & \shortstack[l]{\textbf{Shallow mean diffusive} \\ \textbf{flux magnitude (mg/(mm\textasciicircum{}2 s))}} & \shortstack[l]{\textbf{Mid mean diffusive} \\ \textbf{flux magnitude (mg/(mm\textasciicircum{}2 s))}} & \shortstack[l]{\textbf{Deep mean diffusive} \\ \textbf{flux magnitude (mg/(mm\textasciicircum{}2 s))}} & \shortstack[l]{\textbf{Overall mean diffusive} \\ \textbf{flux magnitude (mg/(mm\textasciicircum{}2 s))}} & \shortstack[l]{\textbf{Mean fixed-ROI} \\ \textbf{95\% CI width (mg/(mm\textasciicircum{}2 s))}} & \shortstack[l]{\textbf{Shallow mean fixed-ROI} \\ \textbf{relative 95\% CI width (\%)}} & \shortstack[l]{\textbf{Mid mean fixed-ROI} \\ \textbf{relative 95\% CI width (\%)}} & \shortstack[l]{\textbf{Deep mean fixed-ROI} \\ \textbf{relative 95\% CI width (\%)}} & \shortstack[l]{\textbf{Profile fit} \\ \textbf{$R^2$}} \\",
        r"\midrule",
    ]

    for tracer in [VIS_LABEL, GAD_LABEL]:
        lines.append(rf"\multicolumn{{12}}{{l}}{{\textbf{{{latex_escape(tracer)}}}}} \\")
        sub = raw_df[raw_df["Tracer"] == tracer].copy()
        panel_order = {name: i for i, name in enumerate(PANEL_NAMES)}
        sub["_panel_order"] = sub["Panel"].map(panel_order)
        sub = sub.sort_values(["_panel_order", "Actual time (min)"])
        for _, row in sub.iterrows():
            lines.append(
                " & ".join([
                    latex_escape(str(row["Panel"])),
                    fmt_fixed(row["Target time (min)"], 2),
                    fmt_fixed(row["Actual time (min)"], 2),
                    fmt_fixed(row[shallow_col], 4),
                    fmt_fixed(row[mid_col], 4),
                    fmt_fixed(row[deep_col], 4),
                    fmt_fixed(row[overall_col], 4),
                    fmt_fixed(row[ci_width_col], 4),
                    fmt_fixed(row["Shallow mean fixed-ROI relative 95% CI width (%)"], 2),
                    fmt_fixed(row["Mid mean fixed-ROI relative 95% CI width (%)"], 2),
                    fmt_fixed(row["Deep mean fixed-ROI relative 95% CI width (%)"], 2),
                    fmt_fixed(row["Profile fit R²"], 3),
                ]) + r" \\")
        if tracer != GAD_LABEL:
            lines.append(r"\midrule")

    lines.extend([
        r"\bottomrule",
        r"\end{tabular}%",
        r"}",
        r"\end{table}",
    ])
    return "\n".join(lines)


def write_latex_document(raw_df: pd.DataFrame, out_path: Path):
    caption = r"\textbf{Table 1.} Early-, mid-, and late-time diffusive flux magnitude summaries for VIS 320 and GAD across shallow, mid, and deep depth bands, with propagated fixed-ROI absolute and relative flux uncertainty."
    label = "tab:early_mid_late_flux_support"
    content = [
        r"\documentclass{article}",
        r"\usepackage{booktabs}",
        r"\usepackage[margin=0.5in]{geometry}",
        r"\usepackage{pdflscape}",
        r"\usepackage{graphicx}",
        r"\usepackage{caption}",
        r"\captionsetup[table]{labelformat=empty,justification=raggedright,singlelinecheck=false}",
        r"\begin{document}",
        r"\begin{landscape}",
        df_to_latex_booktabs_stacked(raw_df, caption, label),
        r"\end{landscape}",
        r"\end{document}",
    ]
    out_path.write_text("\n".join(content), encoding="utf-8")


def write_docx_table(df: pd.DataFrame, out_path: Path) -> bool:
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.shared import Inches
    except Exception:
        print("DOCX skipped: python-docx is not installed.")
        return False

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    doc.add_heading("Early / Mid / Late Flux Support Table", level=1)
    doc.add_paragraph(
        "Table 1. Early-, mid-, and late-time diffusive flux magnitude summaries for VIS 320 and GAD "
        "across shallow, mid, and deep depth bands, with propagated fixed-ROI absolute and relative flux uncertainty."
    )

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr[j].text = str(col)

    for _, row in df.iterrows():
        first_col = str(row.iloc[0])
        if first_col in [VIS_LABEL, GAD_LABEL]:
            cells = table.add_row().cells
            cells[0].text = first_col
            for j in range(1, len(df.columns)):
                cells[j].text = ""
        else:
            cells = table.add_row().cells
            for j, value in enumerate(row.tolist()):
                cells[j].text = str(value)

    doc.save(out_path)
    return True


# ============================================================
# MAIN
# ============================================================
def main():
    ensure_dir(OUT_FOLDER)

    vis = load_tracer_payload(VIS_SOURCE, VIS_RUN_REL, VIS_ROI_FOLDER, VIS_LABEL, VIS_DX_MM)
    gad = load_tracer_payload(GAD_SOURCE, GAD_RUN_REL, GAD_ROI_FOLDER, GAD_LABEL, GAD_DX_MM)

    raw_table = build_support_rows(vis, gad)
    display_table = build_display_df(raw_table)

    out_base = Path(OUT_FOLDER) / "Table_flux_support_early_mid_late_v1"
    save_tables_plain(display_table, out_base)

    if WRITE_LATEX:
        write_latex_document(raw_table, out_base.with_suffix(".tex"))

    wrote_docx = False
    if WRITE_DOCX:
        wrote_docx = write_docx_table(display_table, out_base.with_suffix(".docx"))

    audit = {
        "vis_source": VIS_SOURCE,
        "gad_source": GAD_SOURCE,
        "vis_run_prefix": vis["run_prefix"],
        "gad_run_prefix": gad["run_prefix"],
        "manual_target_times_min": MANUAL_TARGET_TIMES_MIN,
        "panel_names": PANEL_NAMES,
        "flux_uncertainty_method": "Propagated fixed-ROI 95% CI from fitted-profile fixed-ROI uncertainty and fixed-ROI effective-diffusivity uncertainty using low/high bound combinations in J = |-D dC/dx|.",
        "wrote_docx": wrote_docx,
        "wrote_latex": WRITE_LATEX,
        "wrote_markdown": WRITE_MARKDOWN,
        "wrote_csv": WRITE_CSV,
        "wrote_html": WRITE_HTML,
        "outputs": [
            str(out_base.with_suffix(".csv")) if WRITE_CSV else None,
            str(out_base.with_suffix(".md")) if WRITE_MARKDOWN else None,
            str(out_base.with_suffix(".html")) if WRITE_HTML else None,
            str(out_base.with_suffix(".tex")) if WRITE_LATEX else None,
            str(out_base.with_suffix(".docx")) if wrote_docx else None,
        ],
    }
    (Path(OUT_FOLDER) / "Table_flux_support_early_mid_late_v1_audit.json").write_text(json.dumps(audit, indent=2), encoding="utf-8")

    print("Saved early/mid/late flux support table to:", OUT_FOLDER)
    for item in audit["outputs"]:
        if item is not None:
            print(" -", item)


if __name__ == "__main__":
    main()
